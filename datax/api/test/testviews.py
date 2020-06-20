from rest_framework import permissions
from rest_framework.decorators import api_view, permission_classes
from api.utils import *
from connect.olap import OlapClass, kpiTransform

from connect.models import *
from rest_framework.response import Response
import datetime, json
from sqlalchemy import *
from sqlalchemy.sql import select
from sqlalchemy.schema import *
from connect.lib.data import *
from common.head import LOG_COUNT
import random,time
from common import tools
from api import utils

import logging
from common.constantcode import LoggerCode
logger = logging.getLogger(LoggerCode.DJANGOINFO.value)


@api_view(http_method_names=['GET'])
@permission_classes((permissions.AllowAny,))
def test2(request):
    tbs = sourcedetail.objects.get(id=431)
    dist = {}
    dist['rs'] = tbs.distconfig == '[1]'
    return Response(dist)

def test(request):
    from connect.models import olap
    from account.models import sys_userextension, user_tag

    olapRow = olap.objects.get(id=17)
    userRow = sys_userextension.objects.get(id=2)
    userRow.tagfield = [{
        "tagid":"1",
        "value":'城市环境',
        "fun":'like'
    }]

    userRow.save()

    olapRow.tag_config = [{"tagid": "1", "columns": "PR_Contract__DeptName"}]

    olapRow.save()

    result = {
        "aaa":'123'
    }
    return Response(result)

@api_view(http_method_names=['GET'])
@permission_classes((permissions.AllowAny,))
def testprintlogger(request):
    printCount = request.GET['count'] if 'count' in request.GET else 0
    logger.debug('start test logs output! count=%s' % printCount)
    try:
        raise Exception('errrrrrror!')
    except Exception as e:
        # infoLogger = logging.getLogger('djangoInfo')
        # infoLogger.info('------exception---info----')
        # infoLogger.error('------exception--error-----')
        # infoLogger.debug('------exception--debug-----')
        print(e)
    for cnt in range(int(printCount)):
        print('cnt=%s' % cnt)
        logger.error('here is the test log words!here is the test log words!here is the test log words!')
    logger.info('end the test logs output!')
    return Response()

@api_view(http_method_names=['GET'])
@permission_classes((permissions.AllowAny,))
def updateSomeTablesData(request):
    session = SqlUtils(DBTypeCode.EXTENSION_DB.value)
    updateSql = """
    INSERT INTO "public"."test_pollution_0212"("hb_wasteair_control__year","hb_wasteair_control__pollute_quota","hb_wasteair_control__licence_requirement","hb_wasteair_control__environmental_assessment_requirement","hb_wasteair_control__licence_emission","hb_wasteair_control__environmental_assessment_emission","hb_wasteair_control__actual_emission","hb_wasteair_drain__air_volume","add_time","version","extra_processing","keyorder") VALUES 
    """
    t = (1976,1,1,0,0,0,0,0,0)
    e = (2080,12,31,23,59,59,0,0,0)
    starTime = time.mktime(t)
    endTime = time.mktime(e)
    valueLists = []
    for idx in range(1000):
        tempV = []
        tempV.append('\'' + str(2010 + idx % 12) + '\'')
        tempV.append('\'' + str(random.randint(0,10)) + '\'')
        tempV.append('\'' + str(random.randint(0,2)) + '\'')
        tempV.append('\'' + str(random.randint(0,2)) + '\'')
        tempV.append('\'' + str(int(random.uniform(1,3000) * 100) /100) + '\'')
        tempV.append('\'' + str(int(random.uniform(1,3000) * 100) /100) + '\'')
        tempV.append('\'' + str(int(random.uniform(1,3000) * 100) /100) + '\'')
        tempV.append('\'' + str(int(random.uniform(1,3000))) + '\'')
        rndTime = random.randint(starTime, endTime)  # 在开始和结束时间戳中随机取出一个
        tempV.append('\'' + time.strftime("%Y-%m-%d", time.localtime(rndTime)) + '\'')
        tempV.append('\'' + str(random.randint(0,10)) + '\'')
        tempV.append('\'n\'')
        tempV.append('\'' + str(idx) + '\'')
        valueLists.append('(' + ','.join(tempV) + ')')
    updateSql += ','.join(valueLists)
    print('updateSql = ',updateSql)
    session.executeUpdateSql(updateSql)
    session.closeConnect()

    return Response()

#拉取数据库的表信息
@api_view(http_method_names=['GET'])
@permission_classes((permissions.AllowAny,))
def createDBDocument(request):
    # 写word文档文件

    import os
    from docx import Document
    from docx.shared import Inches
    # reload(sys)
    # sys.setdefaultencoding('utf-8')

    # 创建文档对象
    document = Document()

    # 设置文档标题，中文要用unicode字符串
    document.add_heading(u'数据库表信息', 0)

    # 往文档中添加段落
    p = document.add_paragraph('这里获取数据库表的表名、表字段、类型等信息')
    # p.add_run('bold ').bold = True
    # p.add_run('and some ')
    # p.add_run('italic.').italic = True

    # 添加一级标题
    # document.add_heading(u'一级标题, level = 1', level=1)
    # document.add_paragraph('Intense quote', style='IntenseQuote')

    # 添加无序列表
    # document.add_paragraph('first item in unordered list', style='ListBullet')

    # 添加有序列表
    # document.add_paragraph('first item in ordered list', style='ListNumber')
    # document.add_paragraph('second item in ordered list', style='ListNumber')
    # document.add_paragraph('third item in ordered list', style='ListNumber')

    # 添加图片，并指定宽度
    # document.add_picture('e:/docs/pic.png', width=Inches(1.25))

    session = SqlUtils()#获取datax库里所有的表

    getAllTableName = """
        SELECT
            TABLE_NAME,
            COLUMN_NAME,
            data_type
        FROM
            information_schema.COLUMNS
        WHERE
            table_schema = 'public'"""
    querySet = session.getDictResult(getAllTableName)

    print('---querySet=',querySet)
    # 添加表格: 1行3列

    table = document.add_table(rows=1, cols=3)
    tbCell = table.rows[0].cells
    # tbCell[0].text = '表名'
    tbCell[0].text = '字段名'
    tbCell[1].text = '字段类型'

    oldTableName = ''#保存上一个表名
    for singleObj in querySet:
        if oldTableName and oldTableName != singleObj['table_name']:#如果换了表名则重新创建一个表格
            document.add_page_break()

            document.add_paragraph('表' + singleObj['table_name'])#tablename
            table = document.add_table(rows=1, cols=3)

            tbCell = table.rows[0].cells
            # tbCell[0].text = '表名'
            tbCell[0].text = '字段名'
            tbCell[1].text = '字段类型'
            tbCell[2].text = ''

        new_cell = table.add_row().cells
        # new_cell[0].text = singleObj['table_name']
        new_cell[0].text = singleObj['column_name']
        new_cell[1].text = singleObj['data_type']
        new_cell[2].text = ''

        oldTableName = singleObj['table_name']



    # 获取第一行的单元格列表对象
    # hdr_cells = table.rows[0].cells
    # # 为每一个单元格赋值
    # # 注：值都要为字符串类型
    # hdr_cells[0].text = 'Name'
    # hdr_cells[1].text = 'Age'
    # hdr_cells[2].text = 'Tel'
    # # 为表格添加一行
    # new_cells = table.add_row().cells
    # new_cells[0].text = 'Tom'
    # new_cells[1].text = '19'
    # new_cells[2].text = '12345678'

    # 添加分页符
    document.add_page_break()

    # 往新的一页中添加段落
    # p = document.add_paragraph('This is a paragraph in new page.')

    # 保存文档
    saveFilePath = 'd:/temp/dbTableInfo.docx'
    if os.path.exists(saveFilePath):
        os.remove(saveFilePath)
    document.save(saveFilePath)

    return Response()

# region 测试方法
@api_view(http_method_names=['GET'])
@permission_classes((permissions.AllowAny,))
def testGetMethod(request):
    resultObj = tools.successMes()
    try:
        parameter1 = request.GET['parameter1']
        resultObj['data'] = parameter1
    except Exception as error:
        logger.error('---error---file:api.test.testviews.py;method:testGetMethod;error=%s' % error)
        resultObj = tools.errorMes(error.args)
    return Response(resultObj)    

@api_view(http_method_names=['POST'])
@permission_classes((permissions.AllowAny,))
def testPostMethod(request):
    resultObj = tools.successMes()
    try:
        parameter1 = request.data['parameter1']
        executeSession = utils.SqlUtils()
        # 分页查询
        chartsArr = executeSession.getDictDataFromServerPagination(10, 1, {'sql':'select * from dashboard_charts', 'order': 'id'})
        scenes = executeSession.getDictResult('select * from dashboard_scenes')
        resultObj['data'] = {
            'scenes': scenes,
            'chartsArr': chartsArr
        }
    except Exception as error:
        logger.error('---error---file:api.test.testviews.py;method:testGetMethod;error=%s' % error)
        resultObj = tools.errorMes(error.args)
    if executeSession:
        executeSession.closeConnect()
    return Response(resultObj)   
#endregion